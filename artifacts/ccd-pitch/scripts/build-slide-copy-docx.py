#!/usr/bin/env python3
"""Build editorial-review Word doc from active CCDesigner Feature Walkthrough slides."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(
    "/Users/rfreich-storer/CCDesignerARTSCOUNCIL/artifacts/ccd/public/ccd-pitch/"
    "CCDesigner-Feature-Walkthrough-Slide-Copy.docx"
)


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_para(doc, text, *, style=None, size=11, bold=False, italic=False, space_after=6, space_before=0):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_label(doc, label, text, space_after=4):
    """Inline label + copy, e.g. Eyebrow: …"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, size=11, bold=True, color=RGBColor(0x55, 0x55, 0x55))
    r2 = p.add_run(text)
    set_run_font(r2, size=11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_heading1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        set_run_font(run, size=16, bold=True, color=RGBColor(0x00, 0x2D, 0x24))
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_subhead(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True)
    return p


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    # Title
    title = doc.add_heading(
        "CCDesigner Feature Walkthrough — Slide Copy for Editorial Review",
        level=0,
    )
    for run in title.runs:
        set_run_font(run, size=22, bold=True, color=RGBColor(0x00, 0x2D, 0x24))

    add_para(
        doc,
        "Verbatim on-slide copy from all 25 ACTIVE slides in artifacts/ccd-pitch "
        "(slides-manifest.json). Archive slides excluded. Includes eyebrows, headlines, "
        "body, bullets, quotes, stats, CTAs, and meaningful mock UI labels.",
        size=10,
        italic=True,
        space_after=12,
    )
    add_para(
        doc,
        "Source: artifacts/ccd-pitch/src/data/slides-manifest.json + slide components.",
        size=10,
        space_after=6,
    )
    add_hr(doc)

    # ── Slide 1 ──
    add_heading1(doc, "Slide 1 — Exceptional lessons start with connection")
    add_label(doc, "Headline", "Exceptional lessons start with connection")
    add_label(
        doc,
        "Body",
        "Capture ideas. Build lessons. Connect with the best arts organisations — EYFS to A-level.",
    )
    add_para(doc, "Mock UI / brand: CCD (logo mark)", size=10, italic=True, space_after=6)

    # ── Slide 2 ──
    add_heading1(doc, "Slide 2 — A living home for creative teaching")
    add_label(doc, "Eyebrow", "A proposal to the Arts Council")
    add_label(doc, "Headline", "Bringing the arts together — for every child.")
    add_label(
        doc,
        "Body",
        "One connected home for schools, arts organisations and practitioners — "
        "so exceptional creative opportunities reach the children who need them.",
    )
    add_subhead(doc, "Ecosystem labels")
    add_bullets(doc, ["Schools", "Arts orgs", "Practitioners", "Children"])
    add_label(doc, "Footer", "Creative Curriculum Designer")
    add_label(doc, "Footer", "www.ccdesigner.co.uk")

    # ── Slide 3 ──
    add_heading1(doc, "Slide 3 — The problem: arts teachers are working alone")
    add_label(doc, "Eyebrow", "The problem")
    add_label(doc, "Headline", "Arts teachers are working alone.")
    add_label(
        doc,
        "Body",
        "Drama, dance and music teachers stitch a year together from PDFs, shared drives, "
        "photocopied schemes and last year's tabs — often as the only specialist in the building. "
        "Great practice gets lost; non-specialists are left guessing.",
    )
    add_subhead(doc, "Stat")
    add_label(doc, "Figure", "5+")
    add_label(
        doc,
        "Caption",
        "hours a week the average teacher spends planning outside of class — most of it "
        "reinventing what someone, somewhere has already made.",
    )
    add_subhead(doc, "Problem cards")
    add_para(doc, "01 — Practice lives in too many places", bold=True, space_after=2)
    add_para(
        doc,
        "Schemes in folders, rehearsal notes on a phone, warm-ups in a notebook. "
        "Nothing connects, nothing carries forward.",
        space_after=6,
    )
    add_para(doc, "02 — The bigger arc is invisible", bold=True, space_after=2)
    add_para(
        doc,
        "It's hard to see how a term, a key stage, or a whole creative journey fits together — "
        "let alone shape it without starting over.",
        space_after=6,
    )
    add_para(doc, "03 — Creative practice gets lost", bold=True, space_after=2)
    add_para(
        doc,
        "Teaching is creative work — but the warm-up that worked, the rehearsal that landed, "
        "the unit a colleague nailed disappear into forgotten documents by July, learned again "
        "the hard way next September.",
        space_after=6,
    )

    # ── Slide 4 ──
    add_heading1(doc, "Slide 4 — Why now: the arts in higher education are at a tipping point")
    add_label(doc, "Eyebrow", "Why now · the case for the arts")
    add_label(doc, "Headline", "The arts in higher education are at a tipping point.")
    add_label(
        doc,
        "Body",
        "Drama, music and dance courses at UK universities are being merged, paused and closed. "
        "The pipeline that feeds the West End, our orchestras, our touring dance companies — "
        "and every primary classroom that sings — starts in the schools we teach in today. "
        "If the practice we share now is thin, the cohort that arrives at university in five years will be too.",
    )
    add_label(doc, "Quote", '"We don\'t get a second chance at a generation of artists."')
    add_subhead(doc, "Stats")
    add_para(doc, "49", bold=True, space_after=2)
    add_para(
        doc,
        "UK universities have closed, merged or paused arts & humanities courses since 2023.",
        space_after=6,
    )
    add_para(doc, "-47%", bold=True, space_after=2)
    add_para(
        doc,
        "fall in arts GCSE entries since the EBacc was introduced.",
        space_after=6,
    )
    add_para(doc, "1 in 8", bold=True, space_after=2)
    add_para(
        doc,
        "state primaries have no specialist music or drama provision at all.",
        space_after=6,
    )
    add_para(doc, "£126b", bold=True, space_after=2)
    add_para(
        doc,
        "the UK creative industries — built on the artists our classrooms shape.",
        space_after=6,
    )

    # ── Slide 5 ──
    add_heading1(doc, "Slide 5 — The solution: one living home")
    add_label(doc, "Eyebrow", "The solution")
    add_label(doc, "Headline", "One living home for drama, dance and music.")
    add_label(
        doc,
        "Body",
        "Creative Curriculum Designer is a living workspace where great teaching ideas live, "
        "evolve and grow — purpose-built for performing arts education from EYFS through to KS5, "
        "and shared with the wider arts sector.",
    )
    add_subhead(doc, "Pillars")
    add_para(doc, "1 — Half-term planner", bold=True, space_after=2)
    add_para(
        doc,
        "See six half-terms at a glance. Drag, reorder, and shape the year before the year begins.",
        space_after=6,
    )
    add_para(doc, "2 — Lesson library", bold=True, space_after=2)
    add_para(
        doc,
        "Build, reuse and share lessons. Tag by year group and unit. Find the right one in seconds.",
        space_after=6,
    )
    add_para(doc, "3 — Activity stacks", bold=True, space_after=2)
    add_para(
        doc,
        "Group warm-ups, mains and plenaries into reusable stacks. Drop them straight into any lesson.",
        space_after=6,
    )

    # ── Slide 6 ──
    add_heading1(doc, "Slide 6 — Who it's for: arts educators across every key stage")
    add_label(doc, "Eyebrow", "Who it's for")
    add_label(doc, "Headline", "Built with arts educators — every key stage, EYFS to KS5.")
    add_subhead(doc, "Audience: Early Years & Primary")
    add_label(doc, "Meta", "EYFS · KS1 · KS2 · Ages 3–11")
    add_label(
        doc,
        "Subhead",
        "Reception leads, generalist teachers and primary arts specialists",
    )
    add_label(
        doc,
        "Body",
        "Plan around continuous provision and adult-led inputs. Give non-specialist teachers "
        "the confidence to lead drama, dance and music — backed by activities they can actually run.",
    )
    add_label(doc, "Footer", "Continuous provision · non-specialist support")
    add_subhead(doc, "Audience: Secondary & Sixth Form")
    add_label(doc, "Meta", "KS3 · KS4 · KS5 · Ages 11–18")
    add_label(
        doc,
        "Subhead",
        "Drama, dance and music subject leads & departments",
    )
    add_label(
        doc,
        "Body",
        "Sequence units across years, key stages and qualifications — GCSE, A-Level, BTEC and beyond. "
        "Build a shared department library of rehearsal techniques, schemes of work and reflective "
        "practice, so a great unit is taught well by every teacher.",
    )
    add_label(doc, "Footer", "GCSE, A-Level & BTEC frameworks · departmental sharing")
    add_subhead(doc, "Audience: Arts Sector Partners")
    add_label(doc, "Meta", "Theatres · Orchestras · Dance companies · Universities")
    add_label(
        doc,
        "Subhead",
        "Outreach teams, practitioners, training providers and university arts faculties",
    )
    add_label(
        doc,
        "Body",
        "Contribute outreach resources, rehearsal techniques, research-informed frameworks and "
        "training pathways into the schools you partner with — strengthening creative practice "
        "from EYFS to KS5.",
    )
    add_label(doc, "Footer", "Outreach · training · sector-wide collaboration")

    # ── Slide 7 ──
    add_heading1(doc, "Slide 7 — Drama, Music & Dance — one workspace")
    add_label(doc, "Eyebrow", "What it covers")
    add_label(doc, "Headline", "Drama. Music. Dance. One workspace.")
    add_label(
        doc,
        "Body",
        "Three disciplines that share the same rhythm of teaching — warm-up, main, plenary — "
        "modelled end-to-end with starter content across every key stage from EYFS to KS5. "
        "Designed so a non-specialist can step in confidently, and a specialist can build on top.",
    )
    add_subhead(doc, "Drama")
    add_label(doc, "Meta", "EYFS · KS1–KS2 · KS3–KS5")
    add_label(
        doc,
        "Body",
        "Storytelling, voice, character & physical theatre — built around ensemble work and "
        "clear performance moments.",
    )
    add_bullets(
        doc,
        [
            "Storytelling Through Drama",
            "Improvisation & Status",
            "Scripted Scene Work",
            "Devising from Stimulus",
        ],
    )
    add_subhead(doc, "Music")
    add_label(doc, "Meta", "EYFS · KS1–KS2 · KS3–KS5")
    add_label(
        doc,
        "Body",
        "Rhythm, pitch, listening and singing — with backing tracks, vocal guides and resource "
        "links attached to every activity.",
    )
    add_bullets(
        doc,
        [
            "Rhythms of West Africa",
            "Singing & Vocal Health",
            "Body Percussion",
            "Composing with Loops",
        ],
    )
    add_subhead(doc, "Dance")
    add_label(doc, "Meta", "EYFS · KS1–KS2 · KS3–KS5")
    add_label(
        doc,
        "Body",
        "Motif, canon, contact and choreography — sequenced so a class can build a performance "
        "piece across a half-term.",
    )
    add_bullets(
        doc,
        [
            "Motif & Canon",
            "Contact & Trust Work",
            "Storm & Weather Studies",
            "Choreographing in Groups",
        ],
    )

    # ── Slide 8 ──
    add_heading1(doc, "Slide 8 — Feature: Lesson library")
    add_label(doc, "Eyebrow", "Feature 02")
    add_label(doc, "Headline", "Every good lesson, kept.")
    add_label(
        doc,
        "Body",
        "A searchable library of lessons your team builds together — and that grows beyond your "
        "team, with contributions from outreach teams, training providers and arts partners across "
        "the sector. Tag by year group, unit and objective. Open the right one when you need it.",
    )
    add_subhead(doc, "Stats")
    add_para(doc, "200+ — starter lessons included", space_after=4)
    add_para(doc, "1-click — duplicate, edit, share", space_after=6)
    add_subhead(doc, "Mock UI — Lesson library")
    add_label(doc, "Panel title", "Lesson library")
    add_label(doc, "Count", "214 lessons")
    add_label(doc, "Search placeholder", "Search by title, year, or unit…")
    add_para(doc, "Card: Drama · Y4 · 45 min — Freeze-frame storytelling", bold=True, space_after=2)
    add_para(
        doc,
        "Three still images that retell a fable using levels, gesture and gaze.",
        space_after=6,
    )
    add_para(doc, "Card: Music · Y3 · 40 min — Rhythms of West Africa", bold=True, space_after=2)
    add_para(
        doc,
        "Active listening, body percussion, then a class call-and-response.",
        space_after=6,
    )
    add_para(doc, "Card: Dance · Y5 · 50 min — Storm motif & canon", bold=True, space_after=2)
    add_para(
        doc,
        "Build a 16-count motif, then layer it as a four-group canon.",
        space_after=6,
    )
    add_para(doc, "Card: Drama · EYFS · 20 min — Animal walks circle game", bold=True, space_after=2)
    add_para(
        doc,
        "Pretend movement and partner observation in a circle.",
        space_after=6,
    )

    # ── Slide 9 ──
    add_heading1(doc, "Slide 9 — Feature: Activity stacks")
    add_label(doc, "Eyebrow", "Feature 03")
    add_label(doc, "Headline", "Stack activities like Lego.")
    add_label(
        doc,
        "Body",
        "Group your favourite warm-ups, mains and plenaries into reusable stacks. Drop a whole "
        "stack into any lesson and adapt as you go — so brilliant practice survives the end of a "
        "project, a funding cycle, or a school year.",
    )
    add_subhead(doc, "Points")
    add_bullets(
        doc,
        [
            "A — Build once, reuse all year.",
            "B — Share stacks across the team.",
            "C — Drag straight into lesson plans.",
        ],
    )
    add_subhead(doc, "Mock UI — Drama · Storytelling (5 items)")
    add_bullets(
        doc,
        [
            "Mirror game warm-up",
            "Story circle & offers",
            "Character hot-seating",
            "Freeze-frame moments",
            "Audience feedback circle",
        ],
    )
    add_subhead(doc, "Mock UI — Music · Rhythm & voice (4 items)")
    add_bullets(
        doc,
        [
            "Body-percussion warm-up",
            "Active listening: West Africa",
            "Call-and-response clapping",
            "Sing & sign reflection",
            "Drop activity here",
        ],
    )

    # ── Slide 10 ──
    add_heading1(doc, "Slide 10 — Feature: Activity cards with web links")
    add_label(doc, "Eyebrow", "Feature 04")
    add_label(doc, "Headline", "Every resource, one tap away.")
    add_label(
        doc,
        "Body",
        "Activity cards carry the things you actually need in the room: a video demo, a backing "
        "track, a vocal guide, a worksheet, a web link. Pupils can see them on the smartboard. "
        "Teachers click straight through.",
    )
    add_subhead(doc, "Resource types")
    add_bullets(
        doc,
        [
            "Video demos & performance clips",
            "Backing tracks & vocal guides",
            "Web links & teaching ideas",
            "Worksheets & printable resources",
        ],
    )
    add_subhead(doc, "Mock UI — Activity card: Freeze-Frame Storytelling")
    add_label(doc, "Tags", "Drama · Y4 · 20 min · Main activity")
    add_label(
        doc,
        "Body",
        "Groups of four create three still images that retell a familiar fable. Encourage clear "
        "use of levels, facial expression, and gesture.",
    )
    add_bullets(
        doc,
        [
            "BBC Bitesize: Drama warm-up games",
            "Teaching Ideas: Freeze-frames",
        ],
    )
    add_subhead(doc, "Mock UI — Activity card: Active Listening: Rhythms of West Africa")
    add_label(doc, "Tags", "Music · Y3 · 15 min · Warm-up")
    add_label(
        doc,
        "Body",
        "Play a short djembe ensemble recording. Pupils tap the underlying pulse, then identify "
        "the call-and-response pattern and shifting layers.",
    )
    add_bullets(
        doc,
        [
            "Video demo",
            "Backing track",
            "BBC Teach: Djembe ensemble",
        ],
    )
    add_subhead(doc, "Mock UI — Activity card: Storm Motif & Canon")
    add_label(doc, "Tags", "Dance · Y5 · 25 min · Main activity")
    add_label(
        doc,
        "Body",
        "Build a 16-count motif inspired by a storm. In four groups, layer the motif as a canon, "
        "two counts apart.",
    )
    add_bullets(
        doc,
        [
            "Choreography demo",
            "Storm music",
            "Counting worksheet",
        ],
    )

    # ── Slide 11 ──
    add_heading1(doc, "Slide 11 — Feature: Activity creation flow")
    add_label(doc, "Eyebrow", "Feature 05")
    add_label(doc, "Headline", "Build an activity in two minutes.")
    add_label(
        doc,
        "Body",
        "A focused form for the things that matter — name, year groups, time, instructions — "
        "with dedicated fields for the YouTube clip, the backing track, the worksheet and the "
        "web link you'd otherwise lose in a folder.",
    )
    add_subhead(doc, "Points")
    add_para(doc, "Tag once, find forever", bold=True, space_after=2)
    add_para(doc, "Subject, year groups and category attach instantly.", space_after=6)
    add_para(doc, "Web links built in", bold=True, space_after=2)
    add_para(
        doc,
        "Six labelled link slots — video, music, backing, vocals, worksheet, web.",
        space_after=6,
    )
    add_para(doc, "Drag straight in", bold=True, space_after=2)
    add_para(doc, "New activity drops into the lesson you're building.", space_after=6)
    add_subhead(doc, "Mock UI — New activity form")
    add_label(doc, "Panel title", "New activity")
    add_label(doc, "Context", "Lesson 1 · Storytelling Through Drama")
    add_label(doc, "Activity name", "Freeze-Frame Storytelling")
    add_label(doc, "Subject", "Drama")
    add_label(doc, "Category", "Main Activity")
    add_label(doc, "Time (min)", "20")
    add_label(doc, "Year groups", "Year 9 · Year 11 · Year 12 · Year 13")
    add_label(
        doc,
        "Instructions",
        "Groups of four create three still images that retell a familiar fable. Use levels, "
        "facial expression, and gesture.",
    )
    add_label(doc, "Section", "Web links & resources — All optional")
    add_label(doc, "Video link", "youtube.com/watch?v=bP-hnB")
    add_label(doc, "Web resource", "teachingideas.co.uk/drama")
    add_label(doc, "Backing track", "spotify.com/track/...")
    add_label(doc, "Worksheet", "+ Add link")
    add_label(doc, "CTAs", "Cancel · Save activity")

    # ── Slide 12 ──
    add_heading1(doc, "Slide 12 — Feature: Lesson planner calendar")
    add_label(doc, "Eyebrow", "Key feature · Calendar")
    add_label(doc, "Headline", "Your school calendar, your timetable.")
    add_label(
        doc,
        "Body",
        "Set term dates to match your school year. Then drag lessons or whole stacks straight "
        "into the slots where you actually teach.",
    )
    add_subhead(doc, "Points")
    add_para(doc, "Term dates you control", bold=True, space_after=2)
    add_para(
        doc,
        "Autumn, spring and summer blocks align to your inset days and holidays.",
        space_after=6,
    )
    add_para(doc, "Drag lessons into slots", bold=True, space_after=2)
    add_para(
        doc,
        "Drop a single lesson from the library onto any teaching day.",
        space_after=6,
    )
    add_para(doc, "Drop whole stacks", bold=True, space_after=2)
    add_para(
        doc,
        "A stack spreads across your timetable days — one lesson per session.",
        space_after=6,
    )
    add_subhead(doc, "Stats / modes")
    add_para(doc, "Month — term overview", space_after=4)
    add_para(doc, "Timetable — slot-by-slot", space_after=6)
    add_subhead(doc, "Mock UI — Calendar")
    add_label(doc, "Header", "October 2025 · Year 4 Drama")
    add_label(doc, "View toggles", "Month · Week · Timetable")
    add_label(doc, "Day headers", "Mon · Tue · Wed · Thu · Fri · Sat · Sun")
    add_label(doc, "Calendar cells", "Autumn 1 · Freeze-frames · Hot-seating · Inset · Devising · Motif build · Drop stack")
    add_label(doc, "Footer", "Term dates synced · Autumn 1 ends 24 Oct")
    add_label(doc, "CTA", "Edit term dates →")
    add_subhead(doc, "Mock UI — Add to Calendar")
    add_label(doc, "Panel title", "Add to Calendar")
    add_label(doc, "Date", "Sunday, October 12, 2025")
    add_label(doc, "Tabs", "Lesson from Library · Stack")
    add_label(doc, "Search", "Search lessons…")
    add_para(doc, "Commedia KS3 unit — Stack · 4 lessons", bold=True, space_after=2)
    add_para(doc, "Spreads across your Mon & Thu timetable slots", space_after=6)
    add_para(doc, "Freeze-frame storytelling — Drama · Year 4 · 45 min", space_after=4)
    add_para(doc, "Rhythms of West Africa — Music · Year 3 · 40 min", space_after=4)
    add_label(doc, "Footer", "4 sessions from start date")
    add_label(doc, "CTA", "Add to Calendar")

    # ── Slide 13 ──
    add_heading1(doc, "Slide 13 — How it works")
    add_label(doc, "Eyebrow", "How it works")
    add_label(doc, "Headline", "From empty term to ready-to-teach.")
    add_subhead(doc, "Steps")
    add_para(doc, "01 — Map the year", bold=True, space_after=2)
    add_para(
        doc,
        "Open the half-term planner and sketch the arc of your year. Themes, units and topic windows.",
        space_after=6,
    )
    add_para(doc, "02 — Pull from the library", bold=True, space_after=2)
    add_para(
        doc,
        "Drop existing lessons into each half-term or duplicate one and tweak it for a new year group.",
        space_after=6,
    )
    add_para(doc, "03 — Stack the activities", bold=True, space_after=2)
    add_para(
        doc,
        "Add warm-ups, mains and plenaries. Save winning combinations as reusable stacks.",
        space_after=6,
    )
    add_para(doc, "04 — Teach, refine & share", bold=True, space_after=2)
    add_para(
        doc,
        "Print, project, or share with the team — and contribute back to a wider creative archive "
        "the whole sector can draw on.",
        space_after=6,
    )
    add_subhead(doc, "Banner")
    add_label(
        doc,
        "Headline",
        "Built for performing arts teaching from EYFS to KS5 — free for educators.",
    )
    add_label(
        doc,
        "Body",
        "Works on a laptop, a smartboard, a studio tablet, or a phone in the rehearsal room.",
    )
    add_label(doc, "Badges", "Web & PWA · Offline-friendly · Print-ready")
    add_label(doc, "Mark", "CD")

    # ── Slide 14 ──
    add_heading1(doc, "Slide 14 — Partner activity packs & marketplace")
    add_label(doc, "Eyebrow", "Partner & teacher resources")
    add_label(doc, "Headline", "Expert content from partners — and a place for yours.")
    add_label(
        doc,
        "Body",
        "Curated packs from leading arts organisations sit alongside your school library. "
        "Teachers can share packs with colleagues — or sell their own lesson plans to others.",
    )
    add_subhead(doc, "Partner activity packs")
    add_para(doc, "Royal Opera House — Included", bold=True, space_after=2)
    add_para(
        doc,
        "Opera, ballet and vocal warm-ups — specialist expertise ready to drop into your timetable.",
        space_after=2,
    )
    add_para(doc, "12 lesson stacks · video links", space_after=6)
    add_para(doc, "National Theatre — Included", bold=True, space_after=2)
    add_para(
        doc,
        "Devising, text work and ensemble drama — classroom-ready units from NT Learning.",
        space_after=2,
    )
    add_para(doc, "GCSE & A-level · exam-aligned", space_after=6)
    add_para(doc, "More partners joining", bold=True, space_after=2)
    add_para(
        doc,
        "Trinity, ABRSM, Rambert and other specialists — their content accessible inside CCD.",
        space_after=2,
    )
    add_label(doc, "Badge", "Growing library")
    add_label(
        doc,
        "Partner logos (alt text)",
        "Arts Council England · Royal Opera House · London Symphony Orchestra · National Theatre · "
        "BBC Ten Pieces · Tate · The National Gallery · Sadler's Wells",
    )
    add_subhead(doc, "Teacher marketplace")
    add_label(doc, "Eyebrow", "Teacher marketplace")
    add_label(doc, "Headline", "Share free. Sell what you've built.")
    add_label(
        doc,
        "Body",
        "Package your lesson stacks into activity packs. Assign them to your department, share "
        "with partner schools, or list them for purchase.",
    )
    add_label(doc, "Listing", "Commedia KS3 pack — £24.99")
    add_label(doc, "Meta", "By Sarah M. · 8 lessons · Drama")
    add_bullets(
        doc,
        [
            "Share packs with your team at no cost",
            "Sell to other schools and keep creating",
            "Partner content always one click away",
        ],
    )

    # ── Slide 15 ──
    add_heading1(doc, "Slide 15 — Music hubs & Partner Hubs")
    add_label(doc, "Eyebrow", "Partner Hubs")
    add_label(doc, "Headline", "Music hubs in one Partner Hubs home.")
    add_label(
        doc,
        "Body",
        "Dashboard → Partner Hubs → Music section. EMS, Tri-Borough and LSO sit together so "
        "teachers find the right organisation in one place.",
    )
    add_label(doc, "Breadcrumb", "Partner Hubs → Music → EMS · Tri-Borough · LSO")
    add_subhead(doc, "Hub cards")
    add_para(doc, "1 — EMS Music Hub — Interactive", bold=True, space_after=2)
    add_para(doc, "Browse resources · Add to CCDesigner", space_after=2)
    add_para(doc, "Open from Partner Hubs →", space_after=6)
    add_para(doc, "2 — Tri-Borough — Links hub", bold=True, space_after=2)
    add_para(doc, "Music education pathways & partner links", space_after=2)
    add_para(doc, "Open from Partner Hubs →", space_after=6)
    add_para(doc, "3 — LSO hub — Interactive", bold=True, space_after=2)
    add_para(doc, "Choose a pack · Add to CCDesigner", space_after=2)
    add_para(doc, "Open from Partner Hubs →", space_after=6)

    # ── Slide 16 ──
    add_heading1(doc, "Slide 16 — LSO hub — Add to CCDesigner")
    add_label(doc, "Eyebrow", "LSO hub")
    add_label(doc, "Headline", "Choose a resource. Add it to CCDesigner.")
    add_label(
        doc,
        "Body",
        "Open the London Symphony Orchestra hub, pick a classroom pack such as How to Build an "
        "Orchestra, then bring those activities straight into your library — or select Partner "
        "Planning for the term.",
    )
    add_subhead(doc, "Steps")
    add_bullets(
        doc,
        [
            "1 — Browse LSO learning resources",
            "2 — Select a pack for your year group",
            "3 — Add to CCDesigner — activities appear in your library",
        ],
    )
    add_subhead(doc, "Mock UI — London Symphony Orchestra hub")
    add_label(doc, "Header", "London Symphony Orchestra")
    add_label(doc, "Subhead", "Partner hub · Music")
    add_label(doc, "Badge", "Interactive")
    add_para(doc, "How to Build an Orchestra", bold=True, space_after=2)
    add_para(doc, "Year 6 · classroom pack", space_after=2)
    add_label(doc, "CTA", "Add to CCDesigner")
    add_para(doc, "Partner planning", bold=True, space_after=2)
    add_para(doc, "Selectable units for your term overview", space_after=2)
    add_label(doc, "CTA", "Select partner planning")
    add_label(
        doc,
        "Status",
        "Activities land in Activity Library · ready for Lesson Builder — ✓ Synced",
    )

    # ── Slide 17 ──
    add_heading1(doc, "Slide 17 — Activity → Lesson → Term")
    add_label(doc, "Eyebrow", "Planning flow")
    add_label(doc, "Headline", "From activity to lesson to term.")
    add_label(
        doc,
        "Body",
        "Partner activities become a named lesson, then sit on your half-term designer — one "
        "continuous path from sector content to the classroom.",
    )
    add_subhead(doc, "Stages")
    add_para(doc, "1 — Activity Library", bold=True, space_after=2)
    add_para(doc, "Partner & school activities in one place", space_after=6)
    add_para(doc, "2 — Lesson Builder", bold=True, space_after=2)
    add_para(doc, "Sequence warm-ups, mains, plenaries", space_after=6)
    add_para(doc, "3 — Lesson Library", bold=True, space_after=2)
    add_para(doc, "Save, name, and reuse the lesson", space_after=6)
    add_para(doc, "4 — Half-term / term", bold=True, space_after=2)
    add_para(doc, "Place it on your planner overview", space_after=6)

    # ── Slide 18 ──
    add_heading1(doc, "Slide 18 — Key dates & Important dates")
    add_label(doc, "Eyebrow", "Key dates")
    add_label(doc, "Headline", "Populate with partner key dates.")
    add_label(
        doc,
        "Body",
        "Pull Important dates from a partner into your calendar — concerts, workshops, and sector "
        "milestones sit alongside your school year.",
    )
    add_subhead(doc, "Mock UI — Populate control")
    add_label(doc, "Label", "Populate with key dates from")
    add_label(doc, "Dropdown", "London Symphony Orchestra")
    add_subhead(doc, "Mock UI — Important dates")
    add_label(doc, "Panel title", "Important dates")
    add_label(doc, "Meta", "Academic year · 2025–26")
    add_bullets(
        doc,
        [
            "12 Oct — LSO Discovery workshop — Partner",
            "3 Nov — Schools concert — Barbican — Key date",
            "18 Jan — Composer visit (primary) — Partner",
            "9 Mar — Spring showcase deadline — School",
        ],
    )

    # ── Slide 19 ──
    add_heading1(doc, "Slide 19 — Calendar week & timetable")
    add_label(doc, "Eyebrow", "Calendar")
    add_label(doc, "Headline", "Your week, your timetable.")
    add_label(
        doc,
        "Body",
        "See partner lessons and school sessions in week or timetable view — then place library "
        "lessons into the slots you actually teach.",
    )
    add_subhead(doc, "Mock UI — Week view")
    add_label(doc, "Header", "Week view · Year 6 Music")
    add_label(doc, "View toggles", "Month · Week · Timetable")
    add_label(doc, "Days", "Mon · Tue · Wed · Thu · Fri")
    add_bullets(
        doc,
        [
            "Mon — Orchestra build",
            "Wed — Rhythm warm-up",
            "Thu — LSO workshop",
            "Fri — Reflect & share",
        ],
    )

    # ── Slide 20 ──
    add_heading1(doc, "Slide 20 — iCompose & We Teach Drama")
    add_label(doc, "Eyebrow", "Paid partners")
    add_label(doc, "Headline", "iCompose — and We Teach Drama.")
    add_label(
        doc,
        "Body",
        "Premium partner products sit in a clear school basket. A small platform fee (10–20%) "
        "helps sustain CCDesigner while partners keep creating classroom-ready packs.",
    )
    add_subhead(doc, "iCompose")
    add_label(doc, "Badge", "Paid hub")
    add_label(
        doc,
        "Body",
        "Composition pathways for classrooms — browse the hub, add getting-started packs, and "
        "check out through the paid basket.",
    )
    add_label(doc, "CTA", "Add to basket →")
    add_subhead(doc, "We Teach Drama")
    add_label(doc, "Badge", "Paid")
    add_label(
        doc,
        "Body",
        "Specialist drama resources for schools — same basket flow, clear licensing for departments.",
    )
    add_label(doc, "Note", "Brief mention · full hub in prototype")
    add_subhead(doc, "Basket")
    add_label(doc, "Eyebrow", "Basket")
    add_label(doc, "Headline", "Paid partner checkout")
    add_label(
        doc,
        "Body",
        "Platform fee 10–20% · partners keep creating · schools get classroom-ready packs",
    )

    # ── Slide 21 ──
    add_heading1(doc, "Slide 21 — Settings & key-stage folders")
    add_label(doc, "Eyebrow", "Settings")
    add_label(doc, "Headline", "Key-stage folders you shape.")
    add_label(
        doc,
        "Body",
        "Organise your activity library by key stage and category. Choose the academic year from "
        "the year dropdown — your school's structure, not a fixed template.",
    )
    add_bullets(
        doc,
        [
            "Key-stage folders for drama, music & dance",
            "Custom categories and resource links",
            "Academic year dropdown for allocation",
        ],
    )
    add_subhead(doc, "Mock UI — Library · key stages")
    add_label(doc, "Panel title", "Library · key stages")
    add_label(doc, "Year dropdown", "Year · 2025–26 ▾")
    add_bullets(
        doc,
        [
            "EYFS — Customisable folder",
            "KS1 — Customisable folder",
            "KS2 — Customisable folder",
            "KS3 — Customisable folder",
            "KS4 — Customisable folder",
            "KS5 — Customisable folder",
        ],
    )

    # ── Slide 22 ──
    add_heading1(doc, "Slide 22 — Organisation hub & share")
    add_label(doc, "Eyebrow", "Organisations")
    add_label(doc, "Headline", "Share practice across your organisation.")
    add_label(
        doc,
        "Body",
        "School and arts-organisation hubs keep lessons, partner packs, and planning visible to "
        "the right people — so great teaching multiplies beyond one classroom.",
    )
    add_subhead(doc, "Cards")
    add_para(doc, "School hub", bold=True, space_after=2)
    add_para(doc, "Department libraries, branding, and shared units", space_after=6)
    add_para(doc, "Partner hub", bold=True, space_after=2)
    add_para(doc, "Sector resources teachers can pull in once", space_after=6)
    add_para(doc, "Share value", bold=True, space_after=2)
    add_para(doc, "Colleagues reuse what works — nothing lost", space_after=6)

    # ── Slide 23 ──
    add_heading1(doc, "Slide 23 — Celebrating learning in the arts")
    add_label(doc, "Eyebrow", "Celebrating the work")
    add_label(doc, "Headline", "Learning in the arts deserves to be seen.")
    add_label(
        doc,
        "Body",
        "Every freeze-frame, every first improvisation, every Year 6 ensemble piece is evidence "
        "that the arts are alive in our schools. CCD turns the daily craft of teaching into a "
        "record worth celebrating — published to school homepages, shared with parents, gathered "
        "into year-end showcases that tell the real story of arts education.",
    )
    add_subhead(doc, "Cards")
    add_para(doc, "On the school homepage — This week in the arts", bold=True, space_after=2)
    add_para(
        doc,
        "Lessons published live so parents and SLT see what's happening on stage and in the "
        "studio — not just on the spreadsheet.",
        space_after=6,
    )
    add_para(doc, "Year-end showcase — From sketchbook to stage", bold=True, space_after=2)
    add_para(
        doc,
        "Pull every lesson, photo and reflection from a unit into a single, shareable timeline. "
        "Frame the journey, not just the performance.",
        space_after=6,
    )
    add_para(doc, "Practitioner portfolio — What I taught this year", bold=True, space_after=2)
    add_para(
        doc,
        "An export every educator can hand to a head, a partner organisation, or take with them "
        "into the next role.",
        space_after=6,
    )
    add_para(
        doc,
        "Sector evidence — What the arts look like in 1,000 classrooms",
        bold=True,
        space_after=2,
    )
    add_para(
        doc,
        "Anonymised, aggregated insight for funders, MATs and arts councils — so the case for "
        "the arts is grounded in real practice.",
        space_after=6,
    )

    # ── Slide 24 ──
    add_heading1(doc, "Slide 24 — Inspiring, life-changing arts education")
    add_label(doc, "Eyebrow / badge", "Inspiring · Life-changing")
    add_label(doc, "Headline", "One inspiring lesson can change a life.")
    add_label(
        doc,
        "Body",
        "Ask any director, choreographer, conductor or composer where it started, and you'll "
        "hear the same answer: a teacher, a moment, a lesson that opened a door. CCD is built so "
        "those moments multiply — so the brilliant lesson a teacher teaches in Stoke this Tuesday "
        "is the lesson a non-specialist in Truro can teach next week.",
    )
    add_subhead(doc, "Quotes")
    add_para(
        doc,
        '"My drama teacher gave me the courage to audition. Without that one lesson I\'d have '
        'never set foot on a stage."',
        italic=True,
        space_after=2,
    )
    add_para(doc, "— RSC company member, 2024", space_after=6)
    add_para(
        doc,
        '"I learned to read music in Year 5 from a peripatetic teacher who came in once a week. '
        'That teacher built my whole career."',
        italic=True,
        space_after=2,
    )
    add_para(doc, "— BBC Symphony player", space_after=6)
    add_label(
        doc,
        "Callout",
        "Every shared lesson is one more door, opened for one more child.",
    )

    # ── Slide 25 ──
    add_heading1(doc, "Slide 25 — Grow your creative practice — together")
    add_label(
        doc,
        "Badge",
        "Try it with your school, department or arts organisation",
    )
    add_label(doc, "Headline", "Grow your creative practice — together.")
    add_label(
        doc,
        "Body",
        "Built with educators and tested in real classrooms. Powered by your creativity. Shared "
        "with theatres, orchestras, dance companies, universities and outreach teams — together, "
        "building the creative archive your sector deserves. Bring CCD into your school, "
        "department or arts organisation for a half-term trial and we'll walk your team through it.",
    )
    add_subhead(doc, "CTAs")
    add_para(doc, "Visit — www.ccdesigner.co.uk", space_after=4)
    add_para(doc, "Email — hello@ccdesigner.co.uk", space_after=4)
    add_para(
        doc,
        "Book a demo — 30-minute walk-through with your team or partners",
        space_after=6,
    )

    add_hr(doc)
    add_para(
        doc,
        "End of document — 25 active slides. Generated from slide component source for editorial review.",
        size=9,
        italic=True,
        space_after=0,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build()
